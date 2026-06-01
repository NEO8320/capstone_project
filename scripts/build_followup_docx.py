# -*- coding: utf-8 -*-
"""
추가 작업 3건 보고서 docx 생성기 (네트워크 접속 / 새로고침·가중셔플 / 모바일 2열).
build_worklog_docx.py 와 동일한 스타일 헬퍼를 사용한다.
"""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KO_FONT = "맑은 고딕"; CODE_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x38, 0x64); GRAY = RGBColor(0x55, 0x55, 0x55)


def kfont(run, font=KO_FONT, size=None, bold=None, color=None):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    if rf.getparent() is None: rpr.append(rf)
    for a in ("w:ascii","w:hAnsi","w:eastAsia","w:cs"): rf.set(qn(a), font)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color


def shade(p, fill="F4F4F4"):
    pPr = p._p.get_or_add_pPr(); s = OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),fill)
    pPr.append(s)


def body(doc, t, size=10.5, bold=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    kfont(p.add_run(t), size=size, bold=bold, color=color); return p


def heading(doc, t, level=1):
    h = doc.add_heading(level=level)
    kfont(h.add_run(t), size={0:24,1:16,2:12.5}.get(level,11), bold=True, color=ACCENT); return h


def bullet(doc, t):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    kfont(p.add_run(t), size=10.5)


def code(doc, c):
    for line in c.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1); shade(p)
        r = p.add_run(line if line else " "); r.font.size = Pt(9); r.font.name = CODE_FONT
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
        if rf.getparent() is None: rpr.append(rf)
        for a in ("w:ascii","w:hAnsi","w:eastAsia","w:cs"): rf.set(qn(a), CODE_FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def render(doc, blocks):
    for kind, c in blocks:
        if kind=="h2": heading(doc, c, 2)
        elif kind=="p": body(doc, c)
        elif kind=="code": code(doc, c)
        elif kind=="bul":
            for it in c: bullet(doc, it)


SECTIONS = [
("1. 모바일에서도 같은 네트워크로 접속 (WiFi·도메인 없이)", [
    ("h2","요구"),
    ("p","교수님 요청: 시연 데스크탑이 랜선 연결인데, 발표장 공용 WiFi나 도메인 없이 청중이 "
         "폰으로 우리 서비스에 접속하게 하고 싶다."),
    ("h2","핵심 개념"),
    ("p","우리 서비스는 데스크탑 안(localhost)에서 돌고, '같은 네트워크'에 있는 폰이 그 IP로 "
         "접속한다. 인터넷 공개가 아니라 폐쇄 로컬 네트워크라 도메인·저작권·보안 노출이 없다."),
    ("h2","방법 A — 데스크탑을 핫스팟으로 (가장 확실, 추천)"),
    ("p","랜선으로 인터넷에 연결된 데스크탑이 스스로 WiFi를 쏘게 만들어 폰을 붙인다. 발표장 "
         "WiFi가 전혀 필요 없다. (데스크탑에 WiFi 어댑터 필요 — 없으면 USB 무선랜 동글 하나)"),
    ("bul",[
        "설정 → 네트워크 및 인터넷 → 모바일 핫스팟 켜기 → 공유 원본 '이더넷' 선택",
        "폰을 그 핫스팟 SSID에 연결",
        "데스크탑 ipconfig로 핫스팟 IP 확인(보통 192.168.137.1)",
        "폰 브라우저에 http://192.168.137.1:5173 접속",
    ]),
    ("h2","방법 B — 데스크탑이 물린 공유기 WiFi 사용"),
    ("p","데스크탑 랜선이 꽂힌 공유기에 WiFi가 있으면, 폰을 그 WiFi에 연결 → 같은 LAN → "
         "http://<데스크탑 이더넷 IP>:5173 접속. 단 학교·공용 공유기는 'AP 격리'로 기기간 통신을 "
         "막을 수 있어, 그때는 방법 A로 우회한다."),
    ("h2","공통 필수 — 방화벽 허용 (관리자 PowerShell)"),
    ("code","netsh advfirewall firewall add rule name=\"NewsCurator-5173\" dir=in action=allow protocol=TCP localport=5173\n"
            "netsh advfirewall firewall add rule name=\"NewsCurator-8000\" dir=in action=allow protocol=TCP localport=8000"),
    ("h2","청중 편의 — QR 코드"),
    ("p","접속 주소 http://<IP>:5173 를 QR 코드로 만들어 발표 화면/슬라이드에 띄우면 청중이 "
         "카메라로 스캔해 바로 들어온다."),
    ("h2","결과·산출물"),
    ("p","상세 절차·트러블슈팅을 docs/NETWORK_ACCESS_GUIDE.md 로 정리했다. 코드 변경은 없고, "
         "발표 환경 세팅 가이드다. (요약: 방법 A 핫스팟 + 방화벽 허용 + QR 코드가 가장 확실)"),
]),
("2. 메인 피드 새로고침 버튼 + 점수 가중 셔플", [
    ("h2","요구"),
    ("p","메인 피드에 새로고침 버튼을 두고, 추천순으로 정렬하되 새로고침할 때마다 다른 기사들이 "
         "노출되게 하고 싶다."),
    ("h2","설계 — 점수 상위군 가중 셔플"),
    ("p","완전 무작위가 아니라 '추천순의 취지(점수 높은 기사가 상위에 올 확률↑)'를 유지하면서 "
         "매번 순서가 달라지게 했다. Efraimidis-Spirakis 가중 무작위 샘플링: 각 기사에 "
         "key = random()^(1/점수) 를 부여해 key 내림차순 정렬한다. 점수가 클수록 큰 key가 나올 "
         "확률이 높아 상위에 배치되지만, 매 호출 순서가 바뀐다."),
    ("code","// SortFilter.jsx\n"
            "export function weightedShuffle(items) {\n"
            "  return items\n"
            "    .map((it) => {\n"
            "      const w = Math.max(0.0001, Number(it?.score) || 0.0001);\n"
            "      return { it, key: Math.pow(Math.random(), 1 / w) };\n"
            "    })\n"
            "    .sort((a, b) => b.key - a.key)\n"
            "    .map((x) => x.it);\n"
            "}"),
    ("h2","적용 규칙"),
    ("bul",[
        "정렬이 '추천순'일 때만 가중 셔플 적용 — 추천 취지 유지하며 다양화.",
        "'최신순/신뢰도순'은 기존 결정적 정렬 유지(새로고침은 단순 재로드).",
        "새로고침 버튼 클릭 → refreshSeed 증가 + 1페이지로 리셋 + 서버 재요청(loadFeed).",
        "useMemo([feed, hiddenUrls, activeCategory, hideRead, activeSort, refreshSeed])로 묶어 "
        "스크롤·상태변경 시 재셔플되지 않고 버튼 누를 때만 새 조합 생성.",
    ]),
    ("code","// Feed.jsx — 새로고침 핸들러\n"
            "const handleRefresh = () => {\n"
            "  setRefreshSeed((s) => s + 1);   // 가중 셔플 재계산 트리거\n"
            "  setCurrentPage(1);\n"
            "  loadFeed();                      // 캐시 만료 시 최신 데이터 반영\n"
            "};"),
    ("h2","검증"),
    ("p","점수 0.9/0.5/0.2 세 기사를 1000회 셔플해 1등 빈도를 측정: 0.9점 547회 > 0.5점 329회 "
         "> 0.2점 124회. 점수 높을수록 상위 확률이 높으면서도 매번 순서가 달라짐을 확인했다. "
         "프론트 빌드 성공. 파일: Feed.jsx, SortFilter.jsx, Feed.css"),
]),
("3. 모바일 메인 피드 2열 레이아웃", [
    ("h2","증상"),
    ("p","폰에서 메인 피드가 1열로 떠 기사 카드가 화면을 가득 채워 지나치게 컸다. 데스크탑과 "
         "비슷한 밀도로 보고 싶다."),
    ("h2","원인"),
    ("p","index.css가 767px 이하에서 card-grid를 grid-template-columns: 1fr(1열)로 강제하고 "
         "있었다(태블릿은 2열, 데스크탑 3열). 폰만 1열이라 카드가 컸다."),
    ("h2","해결 (diff)"),
    ("code","/* 모바일: 1열 → 2열 */\n"
            "@media (max-width: 767px) {\n"
            "-  .card-grid { grid-template-columns: 1fr; }\n"
            "+  .card-grid { grid-template-columns: repeat(2, 1fr); gap: var(--spacing-sm); }\n"
            "}\n"
            "@media (max-width: 480px) {\n"
            "+  .card-grid { grid-template-columns: repeat(2, 1fr); gap: 8px; }\n"
            "}"),
    ("h2","카드 내부 조정 (ArticleCard.css)"),
    ("bul",[
        "2열 좁은 카드에서 제목이 신뢰도 배지와 겹치지 않게: 제목 padding-right 제거 + "
        "margin-top으로 배지 아래로 내림.",
        "480px 이하: 패딩 10px, 제목 13.5px·최대 3줄, 요약 2줄 말줄임, 배지/메타 폰트 축소, "
        "메타(언론사·기자·시간) 줄바꿈 허용.",
    ]),
    ("h2","결과"),
    ("p","폰(375/430px)에서 카드 2개씩, 데스크탑과 비슷한 밀도로 표시된다. 카드가 좁아져도 "
         "제목·요약·배지가 깨지지 않는다. 프론트 빌드 성공. 파일: index.css, ArticleCard.css"),
]),
]


def build(path):
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = KO_FONT; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)

    for _ in range(5): doc.add_paragraph()
    body(doc, "News Curator", size=32, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    body(doc, "추가 작업 보고서 (시연 대비)", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    body(doc, "모바일 접속 · 새로고침(가중 셔플) · 모바일 2열", size=12.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=24)
    body(doc, "2026-06-01 · GitHub 브랜치 chore/docs-and-scripts", size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88,0x88,0x88))
    doc.add_page_break()

    heading(doc, "개요", level=1)
    t = doc.add_table(rows=1, cols=3); t.style="Light Grid Accent 1"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["#","작업","핵심 파일"]):
        t.rows[0].cells[i].text=""; kfont(t.rows[0].cells[i].paragraphs[0].add_run(h), size=9.5, bold=True)
    for row in [["1","모바일 같은-네트워크 접속(핫스팟/공유기)","docs/NETWORK_ACCESS_GUIDE.md"],
                ["2","새로고침 버튼 + 점수 가중 셔플","Feed.jsx, SortFilter.jsx, Feed.css"],
                ["3","모바일 메인 피드 2열","index.css, ArticleCard.css"]]:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; kfont(cells[i].paragraphs[0].add_run(v), size=9)
    doc.add_paragraph()

    for title, blocks in SECTIONS:
        heading(doc, title, level=1)
        render(doc, blocks)
        doc.add_paragraph()

    doc.save(path); return doc


def dump_md(path):
    L = ["# News Curator — 추가 작업 보고서 (시연 대비)", "",
         "> 모바일 접속 · 새로고침(가중 셔플) · 모바일 2열. 2026-06-01.", ""]
    for title, blocks in SECTIONS:
        L.append(f"## {title}\n")
        for kind, c in blocks:
            if kind=="h2": L.append(f"### {c}\n")
            elif kind=="p": L.append(c+"\n")
            elif kind=="code": L.append("```\n"+c+"\n```\n")
            elif kind=="bul":
                for it in c: L.append(f"- {it}")
                L.append("")
    open(path,"w",encoding="utf-8").write("\n".join(L))


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv)>1 else "docs/News_Curator_추가작업_2026-06-01.docx"
    md = sys.argv[2] if len(sys.argv)>2 else "docs/FOLLOWUP_2026-06-01.md"
    d = build(out); dump_md(md)
    print("docx 저장:", out, "| 문단:", len(d.paragraphs))
    print("md   저장:", md)
