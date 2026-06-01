# -*- coding: utf-8 -*-
"""
2026-06-01 종합 작업보고서 생성기 (docx + md 동시 출력).

내용(content) 데이터를 단일 소스로 정의하고, Word(.docx)와 Markdown(.md)을 함께 출력한다.
- Part 1: 오늘의 작업 15개 항목 (배경/증상/원인/해결+코드/검증/결과 소제목 구조)
- Part 2: 중간발표 피드백 6챕터 (질문/당시답변/실제처리/코드/남은과제)
- 한글 폰트(맑은 고딕), 코드블록(Consolas + 회색 음영), 표지/개요표/부록 포함.

사용:
  python build_worklog_docx.py <out.docx> [out.md]
"""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KO_FONT = "맑은 고딕"
CODE_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x55, 0x55, 0x55)

# ============================================================
# 콘텐츠 데이터 (단일 소스)
#   블록 종류: ("h2", 텍스트) 소제목 | ("p", 텍스트) 본문 | ("code", 코드)
#             | ("bul", [항목...]) 불릿
# ============================================================

OVERVIEW = [
    ["1","인프라","Windows 한글 깨짐 근본 해결(PowerShell shim)","*.bat, scripts/*.ps1","3b4062d"],
    ["2","버그",".env 로드 실패(NAVER 키 미인식)","config.py","3b4062d"],
    ["3","개선","부족 우선 크롤링 배분","pipeline.py","3b4062d"],
    ["4","버그","카테고리 오분류(정치→생활·문화)","pipeline.py","4851dea"],
    ["5","기능","게스트 모드(계정 없이 둘러보기)","auth.py, Login.jsx","4851dea"],
    ["6","확인","라이트테마/계정관리 이미 구현","ThemeContext, Settings","4851dea"],
    ["7","개선","반응형/모바일","index.css, Feed.css","4851dea"],
    ["8","버그","AI 요약 깨짐(\\n·영어 머리말)","llm_processor.py","95c0515"],
    ["9","버그","기자명 추출+구독 버튼","pipeline.py, ArticleDetail.jsx","95c0515"],
    ["10","기능","페이지네이션+읽음 숨김","recommendation.py, Feed.jsx, Pagination.jsx","0f61cf5"],
    ["11","운영","DB 백업/복원 스크립트","backup_db.*, restore_db.*","3b4062d"],
    ["12","문서","README 전면 재작성","README.md","3b4062d"],
    ["13","문서","신뢰도 가중치 학술 근거","CREDIBILITY_RATIONALE.md","4bb2048/6e61964"],
    ["14","문서","크롤링 1시간 주기 근거","CRAWL_INTERVAL_RATIONALE.md","4bb2048"],
    ["15","데이터","깨진 요약 334건 정리+재크롤","(DB 작업)","-"],
]

# ── Part 1: 오늘의 작업 ──
PART1 = [
("1. Windows 한글 깨짐 근본 해결 (PowerShell shim 아키텍처)", [
    ("h2","배경"),
    ("p","이전에 setup.bat의 한글 깨짐은 한 번 고쳤으나, 통합 실행 스크립트 run_all.bat에서 "
         "동일 증상이 재발했다. 조원·다른 PC에서 clone 후 바로 실행해야 하고, 추후 빌드 배포로 "
         "수강생이 실시간 사용할 예정이라 '어느 한국어 Windows에서도 100% 안정'이 목표였다."),
    ("h2","증상 (실제 터미널 출력)"),
    ("code","(.venv) PS> .\\run_all.bat\n"
            "'enAI'은(는) 내부 또는 외부 명령…      ← \"OpenAI\"가 \"enAI\"로 잘림\n"
            "'/b'은(는) 내부 또는 외부 명령…         ← \"exit /b 1\"이 토큰 단위로 깨짐\n"
            "'먯꽌'은(는) 내부 또는 외부 명령…        ← \"먼저\"가 cp949로 오해석"),
    ("p","또한 백엔드를 수동 실행할 때도 $env:PYTHONIOENCODING=\"utf-8\" 를 매번 설정하지 않으면 "
         "uvicorn이 UnicodeEncodeError로 죽었다."),
    ("h2","원인 분석"),
    ("p","한국어 Windows의 cmd.exe는 .bat 파일을 한 줄씩 즉시 파싱하지 않고 시스템 코드페이지"
         "(cp949)로 '미리 버퍼링'한 뒤 파싱한다. 파일을 UTF-8로 저장하고 2번째 줄에 chcp 65001을 "
         "넣어도, 그 명령이 적용되기 전에 이미 후속 라인의 한글 UTF-8 바이트가 cp949의 깨진 글자 + "
         "명령으로 토큰화된다. 즉 순수 .bat로는 회피 불가능한 구조적 문제다."),
    ("h2","해결 — 아키텍처 전환"),
    ("p","'.bat = ASCII 전용 얇은 shim' + '.ps1 = 실제 로직·한글(UTF-8 BOM)'으로 분리했다. "
         "PowerShell 5.1은 Windows 10+에 기본 탑재되어 추가 설치가 없고, .NET 파일 API로 .ps1을 "
         "읽어 BOM 기준으로 인코딩을 결정적으로 식별한다(코드페이지에 종속되지 않음)."),
    ("code","@echo off\nREM run_all.bat — ASCII-only shim\nsetlocal\nset \"SCRIPT_DIR=%~dp0\"\n"
            "powershell.exe -NoProfile -ExecutionPolicy Bypass -File \"%SCRIPT_DIR%scripts\\run_all.ps1\"\n"
            "set \"PS_EXIT=%ERRORLEVEL%\"\necho.\npause\nendlocal & exit /b %PS_EXIT%"),
    ("code","# scripts/run_all.ps1 (UTF-8 BOM)\n"
            "try { [Console]::OutputEncoding = [System.Text.UTF8Encoding]::new($false) } catch {}\n"
            "$env:PYTHONIOENCODING = 'utf-8'\n$env:PYTHONUTF8 = '1'\ntry { chcp 65001 > $null } catch {}\n"
            "# 이후 Docker/Ollama/venv 검증, 백엔드·프론트 새 창 기동"),
    ("h2","추가로 잡은 버그 — PYTHONUTF8 공백"),
    ("p","자식 cmd 창을 'cmd /k \"… && set PYTHONUTF8=1 && …\"' 로 띄우면 cmd의 set이 값 끝에 "
         "공백을 붙여 PYTHONUTF8 값이 '1 '이 되고, Python이 'Fatal Python error: invalid "
         "PYTHONUTF8 environment variable value'로 죽었다. 인라인 set을 제거하고 환경변수는 "
         "PowerShell 부모에서 상속하도록 바꿔 해결했다."),
    ("h2","검증"),
    ("code",".bat 첫 3바이트 = 64 101 99 ('@','e','c') · 비-ASCII 0개\n"
            ".ps1 첫 3바이트 = 239 187 191 (EF BB BF = UTF-8 BOM)"),
    ("h2","결과"),
    ("p","수동 환경변수 없이 더블클릭/PowerShell 실행 모두 한글이 정상 출력되고, 백엔드·프론트가 "
         "각 새 창에서 정상 기동한다. 파일: setup.bat, run_all.bat, scripts/setup.ps1, "
         "scripts/run_all.ps1, start_backend.py, backend/app/main.py"),
]),
("2. 백엔드 .env 로드 실패 (NAVER 키 미인식)", [
    ("h2","증상"),
    ("p","backend/.env에 NAVER_CLIENT_ID/SECRET를 정확히 넣었는데도 백엔드가 "
         "'[Pipeline] NAVER_CLIENT_ID/SECRET 미설정 - 크롤링 불가'를 출력하며 실제 크롤링을 "
         "거부했다. 반면 과거에 'cd backend' 후 uvicorn을 직접 실행했을 때는 크롤링이 됐다."),
    ("h2","원인 분석"),
    ("p","pydantic-settings는 env_file을 '현재 작업 디렉터리(cwd) 기준 상대경로'로 해석한다. "
         "그런데 start_backend.py가 cwd를 프로젝트 루트로 chdir하기 때문에, pydantic은 "
         "<루트>/.env를 찾고 실제 파일 <루트>/backend/.env는 못 찾는다. 그 결과 NAVER 키뿐 "
         "아니라 SECRET_KEY 등 모든 설정이 클래스 기본값(빈 문자열)으로 폴백됐다. 과거 'cd backend' "
         "실행이 우연히 작동했던 것이고, start_backend.py 경로는 처음부터 이 버그를 안고 있었다."),
    ("h2","해결 (diff)"),
    ("code","+from pathlib import Path\n"
            "+# cwd 와 무관하게 모듈 자기 위치 기준 절대경로로 .env 로드\n"
            "+_ENV_FILE = Path(__file__).resolve().parent.parent.parent / \".env\"\n\n"
            " class Settings(BaseSettings):\n     model_config = SettingsConfigDict(\n"
            "-        env_file=\".env\",\n+        env_file=str(_ENV_FILE),\n"
            "         env_file_encoding=\"utf-8\", case_sensitive=False,\n     )"),
    ("h2","검증·결과"),
    ("p","cwd를 프로젝트 루트로 둔 채 설정을 출력해 NAVER_ID 길이 20, SECRET 길이 10, "
         "SECRET_KEY 길이 64로 정상 로드됨을 확인했다. 이후 실제 네이버 뉴스 크롤링이 작동. "
         "파일: backend/app/core/config.py"),
]),
("3. 부족 우선 크롤링 배분", [
    ("h2","배경·문제"),
    ("p","카테고리별 DB 기사 수가 크게 불균형했다(정치 편중). 매 크롤링이 모든 카테고리에 같은 양을 "
         "할당하니, 한 번 벌어진 격차가 좁혀지지 않았다."),
    ("h2","해결"),
    ("p","DB의 카테고리별 기사 수를 조회해, 적은 카테고리에 더 많은 수집 예산을 비례 배분한다. "
         "공식은 weight = (최댓값 - 현재수) + 1 (평탄화로 모든 카테고리에 양수 보장)."),
    ("code","def _allocate_crawl_budget(categories, current_counts, base_per_category,\n"
            "                           min_per_category=5, max_per_category=100):\n"
            "    counts = {c: current_counts.get(c, 0) for c in categories}\n"
            "    max_count = max(counts.values()) if counts else 0\n"
            "    weights = {c: (max_count - counts[c]) + 1 for c in categories}\n"
            "    total_weight = sum(weights.values()) or 1\n"
            "    total_budget = base_per_category * len(categories)\n"
            "    return {c: max(min_per_category,\n"
            "                   min(max_per_category,\n"
            "                       int(round(total_budget*weights[c]/total_weight))))\n"
            "            for c in categories}"),
    ("h2","부가 버그"),
    ("p","처음 카테고리별 카운트를 func.count(Article.id)로 짰는데, Article의 PK가 url이라 "
         "'Article has no attribute id' 오류가 났다. func.count()(행 수 카운트)로 수정했다."),
    ("h2","결과"),
    ("p","부족한 카테고리부터 더 많이 수집 → 시간이 지날수록 분포가 자동으로 균형을 맞춘다. "
         "파일: backend/app/services/pipeline.py"),
]),
("4. 카테고리 오분류 (정치 기사가 생활·문화 탭에 섞임)", [
    ("h2","증상"),
    ("p","파이프라인 로그에서 'Sticky 카테고리: GPT 재분류 정치 -> 크롤러 힌트 생활·문화 유지'가 "
         "반복됐고, 실제로 정치 기사가 생활·문화 탭에 박혔다."),
    ("h2","원인 분석"),
    ("p","빈 탭 문제를 고치며 IT·과학·생활·문화를 sticky 카테고리(크롤러의 키워드 분류를 GPT 재분류"
         "보다 우선)로 넣었던 것이 화근이었다. 두 카테고리의 검색 키워드('문화','IT')가 너무 넓어 "
         "정치(문화체육관광부)·경제(문화산업)·사회(문화행사) 기사가 대량 딸려오는데, GPT는 이를 "
         "정확히 정치/경제/사회로 분류한다. 그런데 sticky가 그 정확한 분류를 무시하고 전부 생활·문화로 "
         "되돌렸다. 세계/연예/스포츠는 키워드가 명확('해외/연예인/스포츠')해 sticky가 타당하지만, "
         "IT·과학·생활·문화는 키워드가 넓어 sticky가 독이 됐다."),
    ("h2","해결 (diff)"),
    ("code"," STICKY_CRAWLER_CATEGORIES = frozenset(\n"
            "-    {\"세계\", \"연예\", \"스포츠\", \"IT·과학\", \"생활·문화\"}\n"
            "+    {\"세계\", \"연예\", \"스포츠\"}\n )"),
    ("h2","결과"),
    ("p","IT·과학·생활·문화는 GPT의 정확한 분류를 신뢰하게 되어 정치 기사가 정치 탭으로 들어간다. "
         "빈 탭 문제는 이미 '서버사이드 카테고리 필터링'으로 해결돼 있어 부작용이 없다. "
         "파일: backend/app/services/pipeline.py (STICKY_CRAWLER_CATEGORIES)"),
]),
("5. 게스트 모드 (계정 없이 둘러보기)", [
    ("h2","요구"),
    ("p","시연 때 평가자·청중이 회원가입 없이 즉시 서비스를 체험할 수 있어야 한다."),
    ("h2","해결 — 백엔드"),
    ("p","공용 게스트 계정 1개를 멱등하게(여러 번 호출해도 1개만) 생성하고 일반 로그인과 동일한 "
         "JWT 토큰을 발급하는 POST /api/auth/guest 엔드포인트를 추가했다. 동시 요청으로 중복 "
         "생성되는 경우를 IntegrityError로 잡아 롤백 후 재조회한다."),
    ("code","GUEST_EMAIL = \"guest@newscurator.demo\"\nGUEST_CATEGORIES = [\"정치\",\"경제\",\"IT·과학\"]\n\n"
            "@router.post(\"/guest\", response_model=TokenResponse)\n"
            "async def guest_login(db=Depends(get_db)):\n"
            "    user = await db.get(User, GUEST_EMAIL)\n"
            "    if not user:\n"
            "        # 콜드스타트 벡터(실패 시 랜덤 정규화 폴백) 후 계정 생성\n"
            "        user = User(email=GUEST_EMAIL, name=\"게스트\", ...)\n"
            "        db.add(user)\n"
            "        try: await db.commit()\n"
            "        except IntegrityError:        # 동시요청 멱등 처리\n"
            "            await db.rollback(); user = await db.get(User, GUEST_EMAIL)\n"
            "    return TokenResponse(access_token=..., refresh_token=...)"),
    ("h2","해결 — 프론트"),
    ("p","로그인 화면에 '계정 없이 둘러보기 (게스트)' 버튼을 추가해 /api/auth/guest 호출 → 토큰 "
         "저장 → /feed로 이동하게 했다."),
    ("h2","검증·결과"),
    ("p","엔드포인트 실제 호출로 토큰 발급·피드 조회(추천 50건 당시)·게스트 계정 1개 멱등 유지를 "
         "확인했다. 파일: backend/app/api/auth.py, frontend/src/components/Login.jsx, Register.css"),
]),
("6. 라이트 테마 / 계정 관리 — 이미 구현됨 확인", [
    ("h2","피드백 맥락"),
    ("p","'화이트(라이트) 테마 가능?', '비밀번호 변경 등 계정 관리 기능?'이라는 피드백이 있었다. "
         "당시 답변은 '다크로 통일/추후 확장'이었으나, 코드를 점검해 보니 둘 다 이미 구현돼 있었다."),
    ("h2","확인 결과"),
    ("bul",[
        "라이트 테마: ThemeContext가 다크↔라이트 토글을 제공하고, index.css에 "
        "[data-theme=\"light\"] 변수 세트가 완비, 헤더에 ThemeToggle이 노출돼 있음.",
        "계정 관리: Settings.jsx에 비밀번호 변경(PATCH /users/me/password)·회원 탈퇴"
        "(DELETE /users/me)가 이미 구현됨.",
    ]),
    ("h2","결과"),
    ("p","추가 개발 없이 '이미 지원됨'으로 정리. 발표에서는 '추후 확장'이 아니라 '구현 완료'로 "
         "답할 수 있다. (다음 발표 답변 강화)"),
]),
("7. 반응형/모바일 대응 보강 (교수님 요청)", [
    ("h2","요구"),
    ("p","교수님이 자연어처리 수업에서 '웹뿐 아니라 모바일에서도 접근 가능하면 좋겠다'고 요청했다."),
    ("h2","보강 내용"),
    ("p","헤더 메뉴가 좁은 화면에서 줄바꿈·중앙정렬되도록, 480px 이하에서 패딩·버튼을 축소, 구독 "
         "트랙 카드는 모바일에서 화면 폭의 85%로 줄여 다음 카드가 살짝 보이게(스와이프 유도) 했다. "
         "카테고리 탭은 기존 가로 스크롤을 유지한다."),
    ("code","@media (max-width: 767px) {\n"
            "  .header { flex-direction: column; align-items: stretch; }\n"
            "  .header__nav { flex-wrap: wrap; justify-content: center; }\n}\n"
            "@media (max-width: 480px) {\n  .card-scroll > * { flex: 0 0 85vw; }\n}"),
    ("h2","결과"),
    ("p","폰(375px)·태블릿(768px)에서 레이아웃이 깨지지 않는다. 빌드 성공으로 CSS 문법 검증 완료. "
         "파일: frontend/src/index.css, Feed.css 등"),
]),
("8. AI 요약 깨짐 수정 (\\n 글자·영어 머리말)", [
    ("h2","증상"),
    ("p","일부 기사의 'AI 3줄 요약'에 줄바꿈 대신 \\n 이 글자 그대로 보이고, 'Here is a 3-line "
         "summary of the article:' 같은 영어 머리말이 본문 앞에 섞여 저장됐다."),
    ("h2","원인 분석"),
    ("p","요약 프롬프트에 '각 줄은 줄바꿈(\\n)으로 구분합니다'라는 지시가 있어 Llama가 리터럴 \\n "
         "문자열을 출력에 넣었고, Llama가 종종 붙이는 영어 안내 머리말을 기존 후처리(요약:/Summary: "
         "제거)가 걸러내지 못했다."),
    ("h2","해결"),
    ("p","프롬프트에서 \\n 지시를 제거(‘실제 줄바꿈으로 3줄’, ‘영어 안내문 금지’)하고, _clean_summary() "
         "후처리 함수를 신설했다: 리터럴 \\n→실제 줄바꿈 변환, 영어/한국어 머리말 반복 제거, 마크다운·"
         "번호·불릿 제거, 한국어가 한 글자도 없으면 None을 반환해 description으로 폴백."),
    ("code","def _clean_summary(text):\n"
            "    text = text.replace('\\\\r\\\\n','\\n').replace('\\\\n','\\n')   # 리터럴→실제 줄바꿈\n"
            "    for _ in range(3):                                    # 머리말 반복 제거\n"
            "        for pat in _SUMMARY_PREFIX_PATTERNS: text = pat.sub('', text, count=1)\n"
            "        text = text.strip()\n"
            "    # 줄별 마크다운/번호/불릿 제거 …\n"
            "    if not re.search(r'[가-힣]', result): return None       # 완전 영어 → 폴백\n"
            "    return result"),
    ("h2","검증·결과"),
    ("p","실제 깨진 요약(리터럴 \\n + 영어 머리말)을 입력해 한국어 3줄만 남고 영어/리터럴이 제거됨을 "
         "확인했다. 추가로 'Here is…' 외에 '다음은 ~ 요약입니다' 한국어 안내 머리말도 발견해 패턴에 "
         "추가했다. 파일: backend/app/services/llm_processor.py"),
]),
("9. 기자명 추출 개선 + 기자 구독 버튼 항상 표시", [
    ("h2","증상"),
    ("p","외부 언론사 기사에서 '기자 구독' 버튼이 아예 보이지 않았다(언론사 구독 버튼만 노출)."),
    ("h2","원인 분석"),
    ("p","기자명 추출이 '본문 끝 200자에서 [이름] 기자' 패턴만 봐서 외부 언론사(본문 끝 표기가 다른 "
         "경우)에서 자주 실패했고, 프론트는 기자명이 없으면(article.journalist가 없으면) 버튼을 "
         "조건부로 숨기고 있었다."),
    ("h2","해결 — 백엔드 다단계 추출"),
    ("p","본문 끝 → 이메일 근접 → 본문 앞 → 역순 어순('기자 홍길동') → description 보조 순으로 "
         "추출을 시도하고, '기자회견·기자단·기자실' 등 일반명사 오매칭은 negative lookahead로 차단했다."),
    ("h2","해결 — 프론트 버튼 항상 표시"),
    ("code","{article.journalist ? (\n"
            "  <button className=\"btn-subscribe\">{article.journalist} 기자 구독</button>\n"
            ") : (\n"
            "  <button className=\"btn-subscribe btn-subscribe--disabled\" disabled\n"
            "          title=\"이 기사는 기자 정보가 확인되지 않아 구독할 수 없습니다.\">\n"
            "    기자 정보 없음\n  </button>\n)}"),
    ("h2","검증·결과"),
    ("p","추출 단위 테스트 8/8 통과(정방향·이메일·역순·앞부분·description 폴백 + '기자회견' 등 오매칭 "
         "차단). 기자명이 없어도 버튼이 사라지지 않고 '기자 정보 없음' 비활성 버튼으로 일관 표시된다. "
         "파일: backend/app/services/pipeline.py, frontend/src/pages/ArticleDetail.jsx, ArticleDetail.css"),
]),
("10. 메인 피드 페이지네이션 + 읽은 기사 숨김 토글", [
    ("h2","요구"),
    ("p","피드에서 기사가 전부 노출되지 않는다는 지적 → 하단에 1·2·3…N 페이지 탭. 그리고 이미 읽은 "
         "기사를 제외하고 볼 수 있는 기능."),
    ("h2","설계 판단"),
    ("p","정렬(추천/최신/신뢰도순)이 클라이언트에서 동작하므로, 페이지 간 정렬 일관성을 위해 "
         "페이지네이션도 클라이언트에서 처리하기로 했다. 백엔드는 반환 상한만 키우고(50→300), "
         "프론트가 [읽음 필터 → 정렬 → 15개씩 페이지 분할]을 담당한다."),
    ("code","-MAX_RECOMMENDATION_TRACK = 50\n-COLD_START_LIMIT = 20\n"
            "+MAX_RECOMMENDATION_TRACK = 300\n+COLD_START_LIMIT = 300"),
    ("h2","프론트 — 파이프라인 + 토글"),
    ("code","const PAGE_SIZE = 15;\n"
            "const filtered = hideRead ? items.filter(i => !i.is_read) : items;\n"
            "const sorted = applySortToItems(filtered, activeSort);\n"
            "const totalPages = Math.max(1, Math.ceil(sorted.length / PAGE_SIZE));\n"
            "const paged = sorted.slice((page-1)*PAGE_SIZE, page*PAGE_SIZE);\n"
            "// 카테고리/정렬/hideRead 변경 시 1페이지로 리셋"),
    ("h2","압축 페이저 (Pagination.jsx)"),
    ("p","페이지가 많아도 전부 나열하지 않고 1 … 4 5 [6] 7 8 … 22 형태로 현재±2와 양끝만 보여준다."),
    ("code","function buildPages(current, total) {\n"
            "  if (total <= 7) return [...Array(total)].map((_,i)=>i+1);\n"
            "  const p=[], L=Math.max(2,current-2), R=Math.min(total-1,current+2);\n"
            "  p.push(1); if(L>2) p.push('…');\n"
            "  for(let i=L;i<=R;i++) p.push(i);\n"
            "  if(R<total-1) p.push('…'); p.push(total); return p;\n}"),
    ("h2","검증·결과"),
    ("p","압축 페이저 로직을 5개 시나리오(22페이지 1/6/끝, 소량, 1페이지)로 검증했고, 읽음 토글은 "
         "기본 노출·켜면 is_read 제외로 동작한다. 모바일·다크/라이트 테마 대응. "
         "파일: recommendation.py, Feed.jsx, Pagination.jsx/css, Feed.css"),
]),
("11. DB 백업/복원 스크립트 (USB 이전용)", [
    ("h2","목적·구현"),
    ("p","시연을 다른 PC에서 할 때 크롤링 데이터를 그대로 옮기기 위해 pg_dump 기반 백업/복원 "
         "스크립트를 만들었다. pgvector 임베딩까지 단일 .dump로 백업하고, 대상 PC에서 복원한다. "
         "ASCII shim(.bat) + PowerShell(.ps1) 동일 구조."),
    ("p","파일: backup_db.bat, restore_db.bat, scripts/backup_db.ps1, scripts/restore_db.ps1"),
]),
("12. README 전면 재작성", [
    ("h2","내용"),
    ("p","작동 방법(처음 실행~매번 실행)·기술 스택·3-tier 아키텍처·동작 원리(크롤링/추천/신뢰도)·"
         "API 엔드포인트·데이터 모델·환경변수·운영(백업/복원)·트러블슈팅·FAQ를 처음부터 새로 정리했다. "
         "게스트 모드·페이지네이션·근거 문서 링크도 반영. 파일: README.md"),
]),
("13. 신뢰도 가중치 학술 근거 (CREDIBILITY_RATIONALE.md)", [
    ("h2","배경"),
    ("p","교수님이 '신뢰도 4지표의 비율(문체30/정보밀도25/인용25/출처20)에 학술적 근거가 있는가'를 "
         "물었다. 받은 논문 3편을 직접 읽고(스캔본은 OCR로 추출) 근거 문서를 작성했다."),
    ("h2","핵심 — 정직한 구분"),
    ("bul",[
        "논문이 뒷받침하는 것 = 가중치의 서열·방향: 문체(불편향) > 정보밀도 ≈ 인용 > 출처.",
        "논문이 도출해 주지 않는 것 = 정확한 % 숫자(30/25/25/20 정밀값). 어떤 논문도 못 줌.",
        "→ 현행 비율은 '논문이 지지하는 서열을 반영한 합리적 배분', 정밀 최적값은 라벨 데이터 튜닝 영역.",
    ]),
    ("p","RB-01 문체30%: '자극적 뉴스 기피'(②) + Meyer/Gaziano '불편향이 최대 요인'. "
         "RB-02 정보밀도25%·RB-03 인용25%: '정확성·전문성=전통 저널리즘 가치'(①③). "
         "RB-04 출처20%: '신뢰할 출처로 전환'(②) — 단독 결정력 낮은 보조 지표."),
]),
("14. 크롤링 1시간 주기 근거 (CRAWL_INTERVAL_RATIONALE.md)", [
    ("h2","6가지 근거"),
    ("bul",[
        "API 한도 여유: 네이버 검색 API 일 25,000회 중 8카테고리×24회≈192회/일(1% 미만).",
        "신선도 체감: 속보가 아닌 한 동일 키워드 상위 결과는 분 단위로 거의 안 바뀜.",
        "처리 비용: 기사당 Llama 요약(수 초)+GPT 분류+Ko-SBERT 임베딩 → 잦으면 부하 누적.",
        "캐시 정합성: 피드 Redis 5분 캐시와 조화.",
        "정보 과부하: 과도한 갱신은 오히려 피로 유발(논문②와 연결).",
        "업계 관행: RSS·뉴스 애그리게이터 통상 15분~1시간.",
    ]),
    ("p","메시지: '더 자주 못 해서가 아니라 할 필요가 없어서' 1시간이며, config로 조정 가능하다."),
]),
("15. 데이터 정비 (깨진 요약 정리 + 재크롤)", [
    ("h2","내용"),
    ("p","기존 DB의 깨진 요약 334건(전체 350건의 95%)을 정리하기 위해 articles 테이블을 비우고"
         "(사용자·구독은 유지) 수정된 코드로 재크롤링했다. 자동 시드(AUTO_SEED_ON_EMPTY_DB)를 끄고 "
         "코드에 박혀 있던 임의 샘플 기사 11건도 제거해, 실제 크롤링 기사만 노출되게 했다."),
    ("h2","디버깅 메모"),
    ("p","작업 중 DB 컨테이너가 down -v로 내려가 백엔드가 degraded 상태였던 적이 있어 "
         "docker-compose up -d로 재기동했고, 게스트 진입 후 피드가 401이던 것은 단순 미로그인 "
         "상태였음을 확인했다."),
]),
]

# ── Part 2: 발표 피드백 6챕터 ──
PART2 = [
("F1. 유사도 연산 최적화 방안 (질문: 장석태, 2026-05-23)", [
    ("h2","질문 원문"),
    ("p","서비스 규모가 커져 하루 수만 건 기사·수만 명 사용자가 되면 매번 수만 유저 벡터와 수만 뉴스 "
         "벡터를 1:1 코사인 유사도 연산하는 것은 서버 자원·응답 속도에 문제가 될 텐데 해결 방법이 있나?"),
    ("h2","중간발표 당시 답변"),
    ("p","지금도 전부 비교하지 않는다. 추천 시 최신 200건만 후보로 추려 그 안에서만 계산하고 결과를 "
         "5분 캐싱한다. DB도 pgvector 인덱스로 가까운 것만 빠르게 찾고, 더 커지면 FAISS 같은 전용 "
         "검색엔진을 붙이면 된다."),
    ("h2","실제 처리 현황"),
    ("p","핵심 설계(후보 제한 + Redis 캐싱 + pgvector HNSW 인덱스)는 그대로 유효하다. 다만 오늘 "
         "'페이지네이션으로 전 기사 노출' 요구를 반영하면서 후보 상한을 키웠다: 반환 상한 "
         "MAX_RECOMMENDATION_TRACK 50→300, 후보 조회는 그 4배(candidate_limit = 300×4 = 1,200)."
         " 즉 후보군이 200→1,200으로 늘었다. 여전히 '전수 비교'가 아니라 '최신순 상위 N건 한정 "
         "계산'이며, FAISS 전환은 대규모 시 과제로 남겨둔 상태다."),
    ("code","# recommendation.py\nMAX_RECOMMENDATION_TRACK = 300     # 반환 상한(50→300)\n"
            "...\n\"candidate_limit\": MAX_RECOMMENDATION_TRACK * 4,   # 후보 1,200건\n"
            "# SQL: ORDER BY published_at DESC LIMIT :candidate_limit"),
    ("h2","남은 과제"),
    ("p","수만~수백만 건 규모에서는 (1) 추천 쿼리를 임베딩 거리 정렬로 바꿔 HNSW 인덱스를 실제로 "
         "태우고 (2) FAISS·전용 벡터DB 도입 (3) 사용자 벡터 사전 배치 계산을 검토한다. 현재 규모"
         "(수백 건)에서는 후보 제한 + 캐싱으로 충분하다."),
]),
("F2. Undo 역산의 L2 정규화 한계 (질문: 박재현, 2026-05-20)", [
    ("h2","질문 원문"),
    ("p","Undo 역산 로직에서 L2 정규화를 적용한 후에는 수식으로 정확한 이전 상태를 복구할 수 없는 것 "
         "아닌가? 정규화 자체가 벡터의 크기 정보를 버리는 연산인데 '완벽한 역산'이라는 표현의 근거는?"),
    ("h2","중간발표 당시 답변"),
    ("p","엄밀히 L2 정규화는 크기 정보를 버리므로 수학적으로 완벽한 복원은 불가능하다. 그래서 Undo는 "
         "'완벽한 역산'이 아니라 방향 기준의 근사 복원에 가깝다. 추천이 코사인 유사도(방향 기반)로 "
         "동작하므로 방향만 보존되면 추천 품질에 실질적 영향이 거의 없다. 더 엄밀히 하려면 로그 "
         "재계산 방식도 고려 중이다."),
    ("h2","실제 처리 현황 — 답변과 코드가 일치(이미 구현)"),
    ("p","feedback.py의 _reverse_ema가 정확히 이 방식이다. 비관심 벡터를 EMA 역산 공식으로 되돌린 "
         "뒤 다시 L2 정규화하며, 결과가 비정상(NaN/inf/영벡터)이면 _is_vector_sane로 막아 적용하지 "
         "않는다. 즉 '방향 근사 복원 + 안전 가드'로 구현돼 있어 답변이 과장이 아니다."),
    ("code","# 정방향 EMA:  V_new = α·V_article + (1-α)·V_old  → L2 정규화\n"
            "# 역산(Undo):  V_old ≈ (V_new - α·V_article) / (1-α) → L2 정규화\n"
            "def _reverse_ema(current_vector, article_vector, alpha):\n"
            "    v_recovered = (np.array(current_vector) - alpha*np.array(article_vector)) / (1-alpha)\n"
            "    norm = np.linalg.norm(v_recovered)\n"
            "    if norm > 0: v_recovered = v_recovered / norm   # 방향 보존(크기 정보는 버림)\n"
            "    return v_recovered.tolist()\n"
            "# 적용 전 _is_vector_sane()로 NaN/inf/영벡터 차단"),
    ("h2","남은 과제"),
    ("p","수학적으로 정확한 복원이 필요하면 답변대로 '피드백 로그를 처음부터 재계산'하는 방식이 대안"
         "이다(현재는 비용 대비 효용이 낮아 근사 복원 유지). 코드 변경 없이 현 상태로 발표 방어 가능."),
]),
("F3. 신뢰도 판단 기준 + 화이트 테마 (질문: 이혜원, 2026-05-21)", [
    ("h2","질문 원문"),
    ("p","신뢰도는 누가 어떤 기준으로 판단하나? 또 화이트 테마로도 사용 가능한가?"),
    ("h2","중간발표 당시 답변"),
    ("p","신뢰도는 문체 중립성·정보 밀도(숫자/날짜/단위)·인용구 존재·출처 명시 4가지를 분석해 산출한다. "
         "화이트 테마는 가독성을 위해 다크로 통일했고 추후 확장 검토 예정."),
    ("h2","실제 처리 현황"),
    ("bul",[
        "신뢰도 기준: 오늘 학술 근거 문서(CREDIBILITY_RATIONALE.md)를 만들어 4지표가 어떤 "
        "선행연구(전통 저널리즘 가치=정확성·공정성·전문성)와 대응하는지, 가중치 서열의 근거가 "
        "무엇인지까지 정리했다. '사람이 아니라 규칙 기반 결정론적 계산'임을 명확히 했다.",
        "화이트 테마: 점검 결과 이미 구현돼 있었다(ThemeContext 다크↔라이트 토글 + "
        "index.css [data-theme=light] 변수 + 헤더 ThemeToggle). 즉 '추후'가 아니라 지금 됨.",
    ]),
    ("h2","남은 과제"),
    ("p","신뢰도 점수의 정량 검증(가짜/정상 라벨 데이터로 분포 비교)은 향후 과제로 문서에 명시했다."),
]),
("F4. 계정 정보(비밀번호 변경 등) 추가 기능 (질문: 이혜원, 2026-05-21)", [
    ("h2","질문 원문"),
    ("p","계정 정보(비밀번호 변경 등) 기능도 구현할 예정인가?"),
    ("h2","중간발표 당시 답변"),
    ("p","현재는 기본 회원가입·로그인만 구현했고 비밀번호 변경 등 계정 관리는 추후 확장 예정."),
    ("h2","실제 처리 현황 — 이미 구현됨 확인"),
    ("p","Settings.jsx 점검 결과 비밀번호 변경(PATCH /users/me/password)과 회원 탈퇴"
         "(DELETE /users/me)가 이미 구현돼 동작한다. 관심 카테고리 수정·글자 크기·테마 전환도 "
         "설정 화면에 있다. 따라서 '추후'가 아니라 '구현 완료'로 답할 수 있다."),
    ("h2","남은 과제"),
    ("p","이메일 변경, 소셜 로그인 등은 미구현(필요 시 확장). 비밀번호 변경 시 현재 비밀번호 확인은 "
         "이미 적용돼 있다."),
]),
("F5. 카테고리 정보 통합 관리 테이블 (질문: 정재윤, 2026-05-19)", [
    ("h2","질문 원문"),
    ("p","관심 카테고리와 기사 카테고리 양쪽에 카테고리 요소가 있는데, 이를 통합 관리하는 테이블을 "
         "운영할 계획이 있나?"),
    ("h2","중간발표 당시 답변"),
    ("p","향후 서비스 확장성과 데이터 정합성을 위해 별도 categories 마스터 테이블로 통합 관리할 계획."),
    ("h2","실제 처리 현황 — 미구현(의도적 보류)"),
    ("p","현재 구조: 사용자의 관심 카테고리는 User.interest_categories(문자열 배열, ARRAY), 기사 "
         "카테고리는 Article.category(문자열)로 각각 저장되고, 코드 상수 CategoryEnum과 config.yaml의 "
         "categories 목록으로 8종을 일관 관리한다. 마스터 테이블은 아직 도입하지 않았다 — 시연을 앞둔 "
         "시점에 DB 스키마 마이그레이션은 리스크가 커서 의도적으로 보류했다(영향: 카테고리 추가·변경 "
         "시 코드 상수와 DB 문자열을 함께 맞춰야 함)."),
    ("h2","남은 과제(계획)"),
    ("p","categories 마스터 테이블(id·name·정렬순서 등)을 만들고 Article·User가 FK로 참조하도록 "
         "정규화하면 (1) 카테고리 추가/이름변경이 한 곳에서 끝나고 (2) 오타·표기 불일치를 원천 차단할 "
         "수 있다. 발표 후 안정화 단계에서 진행 예정."),
]),
("F6. 시간 흐름에 따른 관심사 가중치 (질문: 정재윤, 2026-05-19)", [
    ("h2","질문 원문"),
    ("p","interest_vector와 dislike_vector를 실시간 갱신한다고 했는데, 행동 로그가 많아질수록 오래된 "
         "관심사와 최근 관심사의 가중치를 어떻게 다르게 반영할 계획인가?"),
    ("h2","중간발표 당시 답변"),
    ("p","지수 이동 평균(EMA) 기법으로 최근 관심사에 더 높은 가중치를 부여할 계획."),
    ("h2","실제 처리 현황 — 이미 구현됨"),
    ("p","feedback.py에 EMA가 이미 적용돼 있다. 읽음 피드백은 α=0.15, 관심없음은 α=0.10으로 새 "
         "기사에 가중치를 주고 기존 벡터에 (1-α)를 곱한다. EMA 특성상 오래된 피드백일수록 (1-α)ⁿ으로 "
         "기여가 자동 감쇠하므로, '최근 관심사에 더 높은 가중치'가 수식 차원에서 보장된다."),
    ("code","# feedback.py\nALPHA_READ = 0.15      # 읽음: 새 기사에 15% 가중\n"
            "ALPHA_DISLIKE = 0.10   # 관심없음: 새 기사에 10% 가중\n"
            "# V_new = α·V_article + (1-α)·V_old  → 오래된 기여는 (1-α)^n 로 감쇠"),
    ("h2","남은 과제"),
    ("p","α 값 자체는 실험적으로 정한 값이라 데이터로 튜닝할 여지가 있다(예: 사용자별 적응형 α). "
         "현재 고정값으로도 '최근 우선' 동작은 정상."),
]),
]

# ============================================================
# docx 렌더링 헬퍼
# ============================================================
def set_korean_font(run, font=KO_FONT, size=None, bold=None, color=None):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), font)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color


def shade(p, fill="F4F4F4"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def body(doc, text, size=10.5, bold=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    set_korean_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    sizes = {0: 24, 1: 16, 2: 12.5}
    set_korean_font(h.add_run(text), size=sizes.get(level, 11), bold=True, color=ACCENT)
    return h


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    set_korean_font(p.add_run(text), size=10.5)


def code_block(doc, code):
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        shade(p)
        run = p.add_run(line if line else " ")
        run.font.size = Pt(9)
        run.font.name = CODE_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attr), CODE_FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table(doc, rows, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = ""
        set_korean_font(t.rows[0].cells[i].paragraphs[0].add_run(h), size=9.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            set_korean_font(cells[i].paragraphs[0].add_run(str(v)), size=9)
    return t


def render_blocks(doc, blocks):
    for kind, content in blocks:
        if kind == "h2": heading(doc, content, level=2)
        elif kind == "p": body(doc, content)
        elif kind == "code": code_block(doc, content)
        elif kind == "bul":
            for item in content: bullet(doc, item)


def build_docx(path):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = KO_FONT; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)

    # 표지
    for _ in range(5): doc.add_paragraph()
    body(doc, "News Curator", size=32, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    body(doc, "2026-06-01 종합 작업보고서", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    body(doc, "오늘의 작업 상세 + 중간발표 피드백 대응", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=24)
    body(doc, "AI 기반 개인화 뉴스 큐레이팅 서비스", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=2)
    body(doc, "신규 기능 2 · 버그 수정 5 · 피드백 반영 3 · 문서화 4 · 발표 피드백 6챕터", size=10.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=2)
    body(doc, "GitHub 브랜치: chore/docs-and-scripts", size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88,0x88,0x88))
    doc.add_page_break()

    # 0. 개요
    heading(doc, "0. 작업 개요", level=1)
    body(doc, "2026-06-01에 진행한 전체 작업을 분류·핵심 파일·커밋과 함께 정리한다. "
              "이어지는 Part 1은 작업별 상세, Part 2는 중간발표 피드백 대응이다.")
    table(doc, OVERVIEW, ["#","분류","작업","핵심 파일","커밋"])
    doc.add_page_break()

    # Part 1
    heading(doc, "Part 1. 오늘의 작업 (상세)", level=0)
    doc.add_paragraph()
    for title, blocks in PART1:
        heading(doc, title, level=1)
        render_blocks(doc, blocks)
        doc.add_paragraph()
    doc.add_page_break()

    # Part 2
    heading(doc, "Part 2. 중간발표 피드백 대응", level=0)
    body(doc, "중간발표 후 받은 질문(중복 제외 6종)에 대해 '질문 원문 → 당시 답변 → 실제 처리 현황 "
              "→ 관련 코드 → 남은 과제' 순으로 정리한다. 일부는 오늘 보강했고, 일부는 점검 결과 이미 "
              "구현돼 있었으며, 일부는 의도적으로 보류(계획)한 상태다.", after=10)
    for title, blocks in PART2:
        heading(doc, title, level=1)
        render_blocks(doc, blocks)
        doc.add_paragraph()
    doc.add_page_break()

    # 부록
    heading(doc, "부록. 오늘 커밋 목록 (chore/docs-and-scripts)", level=1)
    table(doc, [
        ["3b4062d","(연속)","인코딩 해결 + 크롤러 개선 + README 재작성 + 백업 스크립트"],
        ["4851dea","15:00","게스트 모드 + 반응형 + 신뢰도 근거 초안 + sticky 수정"],
        ["95c0515","15:18","요약 깨짐 + 기자 구독 버튼"],
        ["0f61cf5","15:32","페이지네이션 + 읽은 기사 숨김"],
        ["4bb2048","15:52","신뢰도 학술 근거 + 크롤링 주기 근거"],
        ["6e61964","16:01","신뢰도 가중치 '비율의 근거와 한계'"],
    ], ["커밋","시각","내용"])

    doc.save(path)
    return doc


def dump_md(path):
    lines = ["# News Curator — 2026-06-01 종합 작업보고서 (상세판)", "",
             "> 오늘의 작업 상세 + 중간발표 피드백 6챕터. 브랜치 chore/docs-and-scripts.", ""]
    lines.append("## 0. 작업 개요\n")
    lines.append("| " + " | ".join(["#","분류","작업","핵심 파일","커밋"]) + " |")
    lines.append("|" + "---|"*5)
    for r in OVERVIEW:
        lines.append("| " + " | ".join(str(x) for x in r) + " |")
    lines.append("")

    def emit(part_title, part):
        lines.append(f"# {part_title}\n")
        for title, blocks in part:
            lines.append(f"## {title}\n")
            for kind, content in blocks:
                if kind == "h2": lines.append(f"### {content}\n")
                elif kind == "p": lines.append(content + "\n")
                elif kind == "code": lines.append("```\n" + content + "\n```\n")
                elif kind == "bul":
                    for it in content: lines.append(f"- {it}")
                    lines.append("")
    emit("Part 1. 오늘의 작업 (상세)", PART1)
    emit("Part 2. 중간발표 피드백 대응", PART2)
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


if __name__ == "__main__":
    out_docx = sys.argv[1] if len(sys.argv) > 1 else "docs/News_Curator_작업보고서_2026-06-01.docx"
    out_md = sys.argv[2] if len(sys.argv) > 2 else "docs/WORKLOG_2026-06-01_DETAILED.md"
    d = build_docx(out_docx)
    dump_md(out_md)
    print("docx 저장:", out_docx, "| 문단:", len(d.paragraphs))
    print("md   저장:", out_md)
